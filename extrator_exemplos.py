import os
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from PIL import Image
from google import genai
from google.genai import types

import database as db

IGNORAR_PALAVRAS_CHAVE = [
    "abteilung",
    "ausbildungsnachweis",
    "ausgeführte arbeiten",
    "unterweisung",
    "betrieblicher unterricht",
    "besondere bemerkungen",
    "stammnr",
    "name:",
    "vorname:",
    "datum",
    "unterschrift",
    "vom:",
    "bis:",
    "nr.:",
    "woche",
    "ausbildungsjahr",
]


def desduplicar_string(s: str) -> str:
    """Corrige repetições automáticas que o Word às vezes gera em campos de formulário/XML."""
    s = s.strip()
    mid = len(s) // 2
    if len(s) % 2 == 0 and s[:mid] == s[mid:]:
        return s[:mid].strip()
    return s


def extrair_texto_de_docx(caminho: Path) -> list[str]:
    """
    Extrai parágrafos e frases técnicas de arquivos Word (.docx e .docm).
    Lê a estrutura OpenXML interna para compatibilidade total com macros e formulários.
    """
    textos = []

    # 1. Leitura direta de XML (suporta .docx e .docm sem restrições)
    try:
        with zipfile.ZipFile(caminho) as z:
            if "word/document.xml" in z.namelist():
                root = ET.fromstring(z.read("word/document.xml"))
                ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                for p in root.iter(f"{ns}p"):
                    txts = [t.text for t in p.iter(f"{ns}t") if t.text]
                    full = "".join(txts).strip()
                    full = desduplicar_string(full)
                    low = full.lower()
                    if len(full) > 20 and not any(ig in low for ig in IGNORAR_PALAVRAS_CHAVE):
                        textos.append(full)
    except Exception:
        pass

    # 2. Fallback para python-docx padrão se o XML direto não capturou
    if not textos:
        try:
            doc = Document(caminho)
            for p in doc.paragraphs:
                t = desduplicar_string(p.text.strip())
                low = t.lower()
                if len(t) > 20 and not any(ig in low for ig in IGNORAR_PALAVRAS_CHAVE):
                    textos.append(t)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = desduplicar_string(p.text.strip())
                            low = t.lower()
                            if len(t) > 20 and not any(ig in low for ig in IGNORAR_PALAVRAS_CHAVE):
                                textos.append(t)
        except Exception as e:
            print(f"Aviso ao ler Word {caminho.name}: {e}")

    return list(dict.fromkeys(textos))


def extrair_texto_de_pdf(caminho: Path) -> list[str]:
    """Extrai texto de relatórios em PDF digital."""
    textos = []
    try:
        reader = PdfReader(caminho)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                for linha in t.split("\n"):
                    linha = desduplicar_string(linha.strip())
                    low = linha.lower()
                    if len(linha) > 20 and not any(ig in low for ig in IGNORAR_PALAVRAS_CHAVE):
                        textos.append(linha)
    except Exception as e:
        print(f"Aviso ao ler PDF {caminho.name}: {e}")
    return list(dict.fromkeys(textos))


def extrair_texto_de_txt(caminho: Path) -> list[str]:
    """Extrai frases de arquivos de texto (.txt)."""
    textos = []
    try:
        conteudo = caminho.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            conteudo = caminho.read_text(encoding="latin-1")
        except Exception:
            return []

    for linha in conteudo.splitlines():
        t = linha.strip()
        low = t.lower()
        if len(t) > 20 and not t.startswith("#") and not any(ig in low for ig in IGNORAR_PALAVRAS_CHAVE):
            textos.append(t)

    return list(dict.fromkeys(textos))


def extrair_texto_de_imagem(caminho: Path, client: genai.Client = None) -> list[str]:
    """
    Extrai apenas o texto contido na imagem.
    Prioriza o cache local em arquivo .txt e no SQLite para NUNCA reenviar a foto desnecessariamente.
    """
    # 1. Verifica se já existe cache em arquivo .txt correspondente
    caminho_txt_cache = caminho.with_suffix(caminho.suffix + ".txt")
    if caminho_txt_cache.exists():
        return extrair_texto_de_txt(caminho_txt_cache)

    caminho_txt_alt = caminho.with_suffix(".txt")
    if caminho_txt_alt.exists() and caminho_txt_alt != caminho:
        return extrair_texto_de_txt(caminho_txt_alt)

    # 2. Verifica se o texto já está gravado no SQLite
    texto_sqlite = db.obter_conteudo_por_arquivo(caminho.name)
    if texto_sqlite:
        linhas = [
            l.strip()
            for l in texto_sqlite.split("\n")
            if len(l.strip()) > 20 and not any(ig in l.lower() for ig in IGNORAR_PALAVRAS_CHAVE)
        ]
        if linhas:
            return linhas

    # 3. Se não houver texto em cache e o cliente Gemini for fornecido, transcreve UMA ÚNICA VEZ
    if client:
        modelos = ["gemini-3.5-flash", "gemini-flash-latest", "gemini-3.7-flash", "gemini-3.5-flash-lite"]
        img = None
        try:
            img = Image.open(caminho)
        except Exception:
            return []

        prompt = (
            "Transcreva exatamente todo o texto, tabelas, horários e matérias desta imagem em alemão. "
            "Retorne o texto das tarefas e matérias de forma limpa e estruturada."
        )

        for mod in modelos:
            try:
                response = client.models.generate_content(
                    model=mod,
                    contents=[img, prompt],
                    config=types.GenerateContentConfig(
                        automatic_function_calling=types.AutomaticFunctionCallingConfig(disable=True)
                    ),
                )
                if response and response.text:
                    texto_transcrito = response.text.strip()
                    # Salva no arquivo de cache .txt para consultas futuras 100% locais
                    try:
                        caminho_txt_cache.write_text(texto_transcrito, encoding="utf-8")
                    except Exception:
                        pass

                    linhas = [
                        l.strip()
                        for l in texto_transcrito.split("\n")
                        if len(l.strip()) > 10 and not any(ig in l.lower() for ig in IGNORAR_PALAVRAS_CHAVE)
                    ]
                    return linhas
            except Exception as e:
                print(f"Aviso ao transcrever imagem {caminho.name} com {mod}: {e}")

    return []


def carregar_base_multimodal(pasta_exemplos: str = "exemplos_antigos", api_key: str = None) -> str:
    """
    Varre os arquivos e carrega EXCLUSIVAMENTE TEXTO para a IA (Gemini).
    Nunca envia imagens brutas no prompt de geração de relatórios.
    """
    pasta = Path(__file__).parent / pasta_exemplos
    if not pasta.exists():
        pasta.mkdir(parents=True, exist_ok=True)
        return ""

    exemplos = []

    # 1. Documentos Word (.docx, .docm)
    arquivos_word = [
        a for a in (list(pasta.glob("*.docx")) + list(pasta.glob("*.docm")))
        if not a.name.startswith("~$")
    ]
    for arq in arquivos_word:
        try:
            exemplos.extend(extrair_texto_de_docx(arq))
        except Exception as e:
            print(f"Erro no Word {arq.name}: {e}")

    # 2. Arquivos PDF (.pdf)
    for arq in pasta.glob("*.pdf"):
        if not arq.name.startswith("~$"):
            exemplos.extend(extrair_texto_de_pdf(arq))

    # 3. Arquivos de Texto (.txt)
    for arq in pasta.glob("*.txt"):
        if not arq.name.startswith("~$") and not arq.name.lower().startswith("readme"):
            exemplos.extend(extrair_texto_de_txt(arq))

    # 4. Textos extraídos de Imagens (lidos apenas do cache de texto / SQLite, sem enviar foto)
    extensoes_img = ("*.png", "*.jpg", "*.jpeg", "*.webp", "*.bmp")
    arquivos_img = []
    for ext in extensoes_img:
        arquivos_img.extend(pasta.glob(ext))

    for arq in arquivos_img:
        if not arq.name.startswith("~$"):
            # Lê apenas o texto extraído correspondente
            textos_img = extrair_texto_de_imagem(arq, client=None)
            exemplos.extend(textos_img)

    # 5. Conhecimento extraído salvo no SQLite
    try:
        textos_sqlite = db.obter_textos_conhecimento()
        for t_sql in textos_sqlite:
            for l in t_sql.splitlines():
                l_str = l.strip()
                if len(l_str) > 20 and not any(ig in l_str.lower() for ig in IGNORAR_PALAVRAS_CHAVE):
                    exemplos.append(l_str)
    except Exception:
        pass

    # Remove duplicatas mantendo a ordem
    exemplos_unicos = list(dict.fromkeys(exemplos))[:25]

    if not exemplos_unicos:
        return ""

    bloco = "\n--- EXEMPLOS DE TEXTO TÉCNICO EXTRAÍDOS ---\n"
    for i, ex in enumerate(exemplos_unicos, 1):
        bloco += f"{i}. {ex}\n"
    bloco += "-------------------------------------------\n"
    return bloco


# Alias para compatibilidade
carregar_base_de_exemplos = carregar_base_multimodal


if __name__ == "__main__":
    api_key_env = os.environ.get("GEMINI_API_KEY", "")
    base = carregar_base_multimodal(api_key=api_key_env)
    print("Texto técnico extraído com sucesso:\n")
    print(base if base else "Nenhum exemplo de texto encontrado.")
